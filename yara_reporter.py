import yaratools
import os
import sys
import textwrap
from datetime import datetime
from docx import Document
import warnings
warnings.filterwarnings("ignore")

def yara_report(rules, results):
    rule_indexes = [i for i, name in enumerate(rules.name) if name in results]
    report_tags = []
    for i in rule_indexes:
        for tag in rules.tags[i]:
            if tag not in report_tags:
                report_tags.extend(rules.tags[i])

    report = []
    for tag in report_tags:
        report_line = {'tag': tag, 'tag_report': []}
        for i in rule_indexes:
            if tag in rules.tags[i]:
                tag_meta = [meta['content'].strip('"') for meta in rules.meta[i]]
                report_line['tag_report'].append(''.join(tag_meta))
        report.append(report_line)
    return report

def listdir_nohidden(path):
	#list directory, excludes hidden files
    for f in os.listdir(path):
        if not f.startswith('.'):
            yield f

def pull_files(d, *args):
	#returns absolute path of all files (including subdirectories) in directory d

	files = []
	for item in listdir_nohidden(d):
		path = os.path.join(d, item)
		if os.path.isfile(path):
			if args:
				extension = os.path.splitext(path)[1]
				if extension in args:
					files.append(os.path.abspath(path))
			else:
				files.append(os.path.abspath(path))
		else:
			files.extend(pull_files(path))
	return files

def main():
    yarafile = sys.argv[1]
    file_arg = sys.argv[2]

    if os.path.isfile(file_arg):
        testfiles = [file_arg]
    else:
        testfiles = pull_files(file_arg)

    with open(yarafile, 'r') as f:
        rawYara = f.read()
    f.close

    rules = yaratools.parse(rawYara)
    print('--------------------------------------')
    document = Document()
    document.add_heading('Yara Report', 0)
    current_time = datetime.now()
    document.add_paragraph(current_time.strftime('Date: %Y-%m-%d'))
    document.add_paragraph(current_time.strftime('Time: %H:%M:%S'))
    document.add_heading('Yara Rules:', 2)
    for name in rules.name:
        document.add_paragraph('%s' % name, style='ListBullet')
    document.add_heading('Files Tested:', 2)
    for testfile in testfiles:
        document.add_paragraph('%s' % os.path.basename(testfile), style='ListBullet')
    document.add_heading('Results:', 2)

    for testfile in testfiles:
        results = rules.runYara(testfile)
        report = yara_report(rules, results)
        print(os.path.basename(testfile))
        document.add_paragraph('%s' % os.path.basename(testfile), style='ListBullet')
        for item in report:
            wrapper = textwrap.TextWrapper(initial_indent='  ', subsequent_indent='  ')
            print(wrapper.fill(item['tag']))
            document.add_paragraph('%s' % item['tag'], style='ListBullet2')
            wrapper = textwrap.TextWrapper(initial_indent='    ', subsequent_indent='    ')
            print(wrapper.fill('\n'.join(item['tag_report'])))
            document.add_paragraph('%s' % '\n'.join(item['tag_report']), style='ListBullet3')
        print('--------------------------------------')
    document.save(current_time.strftime('YARA_report_%Y%m%d_%H%M%S.docx'))

if __name__ == '__main__':
    main()
